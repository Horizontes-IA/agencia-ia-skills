#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generar-accesos.py — Documento de accesos del cliente (Modo Manual de /conectar-cliente).

Lee un accesos.json (ver el esquema en SKILL.md §Modo B) y escribe, en <output_dir>:
    accesos.docx   ← documento Word editable, con la marca de la agencia (si hay python-docx)
    accesos.html   ← documento claro, con la marca de la agencia, imprime a PDF
    accesos.md     ← versión markdown editable

NUNCA guarda claves: el campo `ubicacion` trae un placeholder [...] a propósito.
La marca (nombre, color, contacto) sale de ~/.config/agencia-ia/perfil.json.

Uso:  python3 generar-accesos.py <accesos.json> <output_dir>
"""
import sys, os, json, html
from pathlib import Path


def esc(s):
    return html.escape("" if s is None else str(s), quote=True)


def marca():
    """Marca de la agencia (nombre, color de acento, contacto) desde el perfil."""
    base = {"nombre": "Horizontes IA", "color": "#00E5FF", "contacto": ""}
    try:
        p = json.loads(Path(os.path.expanduser("~/.config/agencia-ia/perfil.json")).read_text(encoding="utf-8"))
    except Exception:
        return base
    ag = p.get("agencia", {}) or {}; mk = p.get("marca", {}) or {}; pe = p.get("persona", {}) or {}
    color = (mk.get("color_acento") or "").strip()
    base["nombre"] = ag.get("nombre_marca") or base["nombre"]
    if len(color) == 7 and color[0] == "#":
        base["color"] = color
    base["contacto"] = (pe.get("web") or pe.get("email") or pe.get("whatsapp") or "").strip()
    return base


def on_accent(hex_color):
    """Blanco si el acento es oscuro, tinta si es claro (contraste)."""
    try:
        r, g, b = int(hex_color[1:3], 16), int(hex_color[3:5], 16), int(hex_color[5:7], 16)
        return "#ffffff" if (0.299 * r + 0.587 * g + 0.114 * b) < 150 else "#15181e"
    except Exception:
        return "#15181e"


def build_html(d, m):
    acc = d.get("accesos", []) or []
    rows = ""
    for a in acc:
        rows += (
            "<tr>"
            f"<td class='app'>{esc(a.get('app'))}</td>"
            f"<td>{esc(a.get('metodo'))}</td>"
            f"<td>{esc(a.get('usuario'))}</td>"
            f"<td class='loc'>{esc(a.get('ubicacion'))}</td>"
            f"<td class='muted'>{esc(a.get('notas'))}</td>"
            "</tr>"
        )
    nota = d.get("nota_seguridad") or ("Este documento NO contiene las claves. Guárdalas en un gestor de "
                                       "contraseñas; nunca por correo o WhatsApp.")
    contacto = f" · {esc(m['contacto'])}" if m["contacto"] else ""
    ACC, OA = m["color"], on_accent(m["color"])
    return f"""<!DOCTYPE html><html lang="es"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Accesos — {esc(d.get('cliente'))} — {esc(m['nombre'])}</title>
<link href="https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>
:root{{--acc:{ACC}; --ink:#15181e; --muted:#5b6573; --line:rgba(17,24,39,.12); --surface:#f7f9fc;}}
*{{box-sizing:border-box;}} body{{margin:0; background:#fff; color:var(--ink);
  font-family:'Space Grotesk',system-ui,sans-serif; font-size:15px; line-height:1.55;}}
.page{{max-width:900px; margin:0 auto; padding:48px 44px;}}
.eyebrow{{color:var(--acc); font-size:11px; font-weight:700; letter-spacing:.18em; text-transform:uppercase;}}
h1{{font-size:30px; font-weight:700; margin:6px 0 2px; letter-spacing:-.01em;}}
.meta{{color:var(--muted); font-size:14px; margin-bottom:6px;}}
.rule{{height:3px; width:56px; background:var(--acc); border-radius:3px; margin:16px 0 26px;}}
table{{width:100%; border-collapse:collapse; margin-top:6px; font-size:14px;}}
th{{text-align:left; background:var(--surface); color:var(--muted); text-transform:uppercase; letter-spacing:.06em;
  font-size:11px; font-weight:700; padding:11px 12px; border-bottom:2px solid var(--acc);}}
td{{padding:12px; border-bottom:1px solid var(--line); vertical-align:top;}}
td.app{{font-weight:600;}} td.loc{{color:var(--acc); font-weight:600;}} td.muted{{color:var(--muted);}}
.warn{{margin-top:28px; border:1px solid var(--line); border-left:4px solid var(--acc);
  background:var(--surface); border-radius:0 10px 10px 0; padding:16px 20px; font-size:13.5px; color:#3d4753;}}
.warn b{{color:var(--ink);}}
footer{{margin-top:34px; padding-top:16px; border-top:1px solid var(--line);
  color:var(--muted); font-size:12px; text-align:center;}}
footer b{{color:var(--acc);}}
@page{{size:A4; margin:14mm;}}
@media print{{*{{-webkit-print-color-adjust:exact !important; print-color-adjust:exact !important;}}
  tr{{break-inside:avoid;}} p,li{{orphans:3; widows:3;}}}}
</style></head><body><div class="page">
<div class="eyebrow">Accesos del proyecto — confidencial</div>
<h1>Accesos de {esc(d.get('cliente'))}</h1>
<div class="meta">Preparado por <b>{esc(m['nombre'])}</b>{contacto} · {esc(d.get('fecha',''))}</div>
<div class="rule"></div>
<table>
<tr><th>Cuenta / App</th><th>Método</th><th>Usuario</th><th>Dónde vive la credencial</th><th>Notas</th></tr>
{rows}
</table>
<div class="warn"><b>⚠ Seguridad:</b> {esc(nota)}</div>
<footer><b>{esc(m['nombre'])}</b> · Documento de accesos · guárdalo en un lugar seguro</footer>
</div></body></html>"""


def build_md(d, m):
    L = [f"# Accesos de {d.get('cliente','')}", "",
         f"**Preparado por:** {m['nombre']}" + (f" · {m['contacto']}" if m["contacto"] else ""),
         f"**Fecha:** {d.get('fecha','')}", "",
         "| Cuenta / App | Método | Usuario | Dónde vive la credencial | Notas |",
         "|---|---|---|---|---|"]
    for a in d.get("accesos", []) or []:
        L.append("| {} | {} | {} | {} | {} |".format(
            a.get("app", ""), a.get("metodo", ""), a.get("usuario", ""),
            a.get("ubicacion", ""), a.get("notas", "")))
    L += ["", "> ⚠ **Seguridad:** " + (d.get("nota_seguridad") or
          "Este documento NO contiene las claves. Guárdalas en un gestor de contraseñas; nunca por correo o WhatsApp.")]
    return "\n".join(L)


def build_docx(d, m, out_path):
    """Documento Word editable, con la marca de la agencia. Requiere python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    c = m["color"].lstrip("#")
    ACC = RGBColor(int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16))
    INK, MUTED = RGBColor(0x15, 0x18, 0x1E), RGBColor(0x5B, 0x65, 0x73)
    doc = Document()
    base = doc.styles["Normal"]; base.font.name = "Calibri"; base.font.size = Pt(11); base.font.color.rgb = INK

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("ACCESOS DEL PROYECTO — CONFIDENCIAL"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = ACC
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Accesos de " + (d.get("cliente") or "")); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = INK
    contacto = (" · " + m["contacto"]) if m["contacto"] else ""
    p = doc.add_paragraph()
    r = p.add_run("Preparado por " + m["nombre"] + contacto + " · " + (d.get("fecha") or "")); r.font.size = Pt(10); r.font.color.rgb = MUTED

    headers = ["Cuenta / App", "Método", "Usuario", "Dónde vive la credencial", "Notas"]
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 2"
    for i, h in enumerate(headers):
        run = t.rows[0].cells[i].paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(9); run.font.color.rgb = ACC
    for a in d.get("accesos", []) or []:
        cells = t.add_row().cells
        for i, v in enumerate([a.get("app"), a.get("metodo"), a.get("usuario"), a.get("ubicacion"), a.get("notas")]):
            cells[i].text = "" if v is None else str(v)

    doc.add_paragraph()
    nota = d.get("nota_seguridad") or ("Este documento NO contiene las claves. Guárdalas en un gestor de "
                                       "contraseñas; nunca por correo o WhatsApp.")
    p = doc.add_paragraph()
    r = p.add_run("⚠ Seguridad:  "); r.bold = True; r.font.color.rgb = ACC
    r2 = p.add_run(nota); r2.font.size = Pt(10.5)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(m["nombre"] + " · Documento de accesos · guárdalo en un lugar seguro"); r.font.size = Pt(9); r.font.color.rgb = MUTED
    doc.save(str(out_path))


def main():
    if len(sys.argv) != 3:
        sys.exit("Uso: python3 generar-accesos.py <accesos.json> <output_dir>")
    try:
        d = json.loads(Path(sys.argv[1]).read_text(encoding="utf-8"))
    except Exception as e:
        sys.exit("No pude leer el accesos.json: {}".format(e))
    out = Path(sys.argv[2]); out.mkdir(parents=True, exist_ok=True)
    m = marca()
    (out / "accesos.html").write_text(build_html(d, m), encoding="utf-8")
    (out / "accesos.md").write_text(build_md(d, m), encoding="utf-8")
    try:
        build_docx(d, m, out / "accesos.docx")
        print("Accesos → {} (+ .docx editable, + .html para PDF)".format(out / "accesos.docx"))
    except ImportError:
        print("Accesos → {} (.html + .md). Para el .docx editable: pip3 install python-docx".format(out / "accesos.html"))


if __name__ == "__main__":
    main()
